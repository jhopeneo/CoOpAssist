"""
Word document loader for QmanAssist.
Uses python-docx to extract text, tables, and structure from .docx files.
Uses docx2txt for legacy .doc file support.
Supports both local files and SMB network shares.
"""

from pathlib import Path
from typing import List
from loguru import logger
import docx
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
import docx2txt
import smbclient
import io

from .base_loader import BaseDocumentLoader, Document


class WordDocumentLoader(BaseDocumentLoader):
    """Loader for Microsoft Word documents (.doc and .docx)."""

    def __init__(self, file_path: Path, preserve_structure: bool = True):
        """Initialize Word document loader.

        Args:
            file_path: Path to .doc or .docx file.
            preserve_structure: Whether to preserve document structure (headings, paragraphs).
                Note: Structure preservation only works for .docx files.
        """
        super().__init__(file_path)
        self.preserve_structure = preserve_structure

    def get_supported_extensions(self) -> List[str]:
        """Get supported file extensions."""
        return [".doc", ".docx", ".DOC", ".DOCX"]

    def load(self) -> List[Document]:
        """Load Word document and extract content.

        Returns:
            List containing a single Document object with full document content.
        """
        try:
            # Check file extension to determine which method to use
            extension = self.file_path.suffix.lower()

            if extension == ".doc":
                # Use docx2txt for legacy .doc files (text extraction only)
                return self._load_doc()
            else:
                # Use python-docx for .docx files (full structure)
                return self._load_docx()

        except Exception as e:
            logger.error(f"Error loading Word document {self.file_path}: {e}")
            raise

    def _load_docx(self) -> List[Document]:
        """Load .docx file with full structure preservation.

        Returns:
            List containing a single Document object.
        """
        try:
            # Check if it's an SMB path
            path_str = str(self.file_path)
            if path_str.startswith("//") or path_str.startswith("\\\\"):
                # Read from SMB into memory
                smb_path = path_str.replace("//", "\\\\").replace("/", "\\")
                with smbclient.open_file(smb_path, mode="rb") as smb_file:
                    docx_data = io.BytesIO(smb_file.read())
                    doc = docx.Document(docx_data)
            else:
                # Local file
                doc = docx.Document(self.file_path)

            base_metadata = self._get_base_metadata()

            # Extract content with structure
            content_parts = []
            current_section = None

            # Track document structure
            heading_levels = []

            for element in doc.element.body:
                # Check if it's a paragraph
                if element.tag.endswith('}p'):
                    para = Paragraph(element, doc)
                    text = para.text.strip()

                    if text:
                        # Check if it's a heading
                        if para.style.name.startswith('Heading'):
                            heading_level = para.style.name
                            heading_levels.append((heading_level, text))
                            content_parts.append(f"\n## {text}\n")
                            current_section = text
                        else:
                            content_parts.append(text)

                # Check if it's a table
                elif element.tag.endswith('}tbl'):
                    table = Table(element, doc)
                    table_text = self._extract_table(table)
                    if table_text:
                        content_parts.append(f"\n{table_text}\n")

            # Combine all content
            content = "\n".join(content_parts)
            content = self._clean_text(content)

            # Update metadata
            base_metadata.update({
                "doc_type": "docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "has_tables": len(doc.tables) > 0,
                "heading_count": len(heading_levels),
            })

            # Add core properties if available
            try:
                core_props = doc.core_properties
                if core_props.title:
                    base_metadata["title"] = core_props.title
                if core_props.author:
                    base_metadata["author"] = core_props.author
                if core_props.subject:
                    base_metadata["subject"] = core_props.subject
            except Exception as e:
                logger.debug(f"Could not extract core properties: {e}")

            documents = [Document(content=content, metadata=base_metadata)]

            logger.info(
                f"Loaded Word document {self.file_path.name}: "
                f"{base_metadata['paragraph_count']} paragraphs, "
                f"{base_metadata['table_count']} tables"
            )

            return documents

        except Exception as e:
            logger.error(f"Error loading .docx document {self.file_path}: {e}")
            raise

    def _load_doc(self) -> List[Document]:
        """Load legacy .doc file using docx2txt (text extraction only).

        Returns:
            List containing a single Document object.
        """
        try:
            base_metadata = self._get_base_metadata()
            base_metadata["doc_type"] = "doc"

            # Check if it's an SMB path
            path_str = str(self.file_path)
            if path_str.startswith("//") or path_str.startswith("\\\\"):
                # Read from SMB into memory, then extract text
                smb_path = path_str.replace("//", "\\\\").replace("/", "\\")
                with smbclient.open_file(smb_path, mode="rb") as smb_file:
                    # docx2txt needs a file path, so we'll write to temp and read
                    import tempfile
                    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
                        tmp.write(smb_file.read())
                        tmp_path = tmp.name

                    try:
                        content = docx2txt.process(tmp_path)
                    finally:
                        # Clean up temp file
                        import os
                        os.unlink(tmp_path)
            else:
                # Local file - direct extraction
                content = docx2txt.process(str(self.file_path))

            # Clean the extracted text
            content = self._clean_text(content)

            if not content or not content.strip():
                logger.warning(f"No text content extracted from {self.file_path.name}")
                content = "[Empty document]"

            documents = [Document(content=content, metadata=base_metadata)]

            logger.info(f"Loaded legacy .doc document {self.file_path.name}")

            return documents

        except Exception as e:
            logger.error(f"Error loading .doc document {self.file_path}: {e}")
            raise

    def _extract_table(self, table: Table) -> str:
        """Extract text from a Word table.

        Args:
            table: python-docx Table object.

        Returns:
            Formatted table text.
        """
        table_lines = []

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)

            if any(row_data):  # Only add non-empty rows
                row_text = " | ".join(row_data)
                table_lines.append(row_text)

        if table_lines:
            return "Table:\n" + "\n".join(table_lines)

        return ""
